"""
硬件项目管理系统 - Flask 后端
读取 总表.xls，合并本地 JSON 商务信息，提供 REST API
"""
import os
import json
import re
import sys
import time
import zipfile
import io
from datetime import datetime
from flask import Flask, jsonify, request, render_template, make_response, send_file
import xlrd
import xlwt
from xlutils.copy import copy as xl_copy
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.config['JSON_ENSURE_ASCII'] = False  # 允许非ASCII字符直接编码

# 路径配置
# 打包为 exe 时，用环境变量 APP_BASE_DIR 指向 exe 所在目录（由 app_main.py 注入）
# 直接运行 server.py 时，用 __file__ 推导
def _get_base_dir():
    env = os.environ.get('APP_BASE_DIR')
    if env:
        return env
    return os.path.dirname(os.path.abspath(__file__))

BASE_DIR = _get_base_dir()           # project_manager/ 目录（含 data/）
# 打包后：sys._MEIPASS（_internal/），总表.xls 在这里
# 直接运行：project_manager/，总表.xls 在父目录（Working Repository/）
parent_dir = os.path.dirname(BASE_DIR)  # 父目录：Working Repository/
# 优先读取 .xlsx，兼容旧的 .xls
XLS_PATH = os.path.join(parent_dir, '总表.xlsx')  # 从父目录读取（新格式）
if not os.path.exists(XLS_PATH):
    # 兼容旧版本
    XLS_PATH = os.path.join(parent_dir, '总表.xls')
if not os.path.exists(XLS_PATH):
    # 如果父目录都没有，尝试项目目录本身
    XLS_PATH = os.path.join(BASE_DIR, '总表.xlsx')
if not os.path.exists(XLS_PATH):
    XLS_PATH = os.path.join(BASE_DIR, '总表.xls')
# 项目文件夹在父目录的 projects/
if os.path.isdir(os.path.join(os.path.dirname(BASE_DIR), 'projects')):
    PROJECTS_DIR = os.path.join(os.path.dirname(BASE_DIR), 'projects')
else:
    PROJECTS_DIR = os.path.join(BASE_DIR, 'projects')
DATA_PATH = os.path.join(BASE_DIR, 'data', 'projects.json')
PARTS_PATH = os.path.join(BASE_DIR, 'data', 'parts.json')  # 器件库
EXPORT_RECORDS_PATH = os.path.join(BASE_DIR, 'data', 'export_records.json')  # 导出记录


# ─── 数据读写 ────────────────────────────────────────────────

def load_json():
    if not os.path.exists(DATA_PATH):
        return {}
    with open(DATA_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_json(data: dict):
    os.makedirs(os.path.dirname(DATA_PATH), exist_ok=True)
    with open(DATA_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_parts():
    """读取器件库 JSON
    结构: { "DHT11": { "stock": 2, "purchases": [{"qty":2,"date":"2024-01-01","note":""}] } }
    """
    if not os.path.exists(PARTS_PATH):
        return {}
    with open(PARTS_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_parts(data: dict):
    os.makedirs(os.path.dirname(PARTS_PATH), exist_ok=True)
    with open(PARTS_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_export_records():
    """读取导出记录 JSON
    结构: { "records": [{ "id": "...", "project_name": "...", "export_time": "...", "files": [...], ... }] }
    """
    if not os.path.exists(EXPORT_RECORDS_PATH):
        return {"records": []}
    try:
        with open(EXPORT_RECORDS_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError):
        return {"records": []}


def save_export_records(data: dict):
    """保存导出记录"""
    os.makedirs(os.path.dirname(EXPORT_RECORDS_PATH), exist_ok=True)
    with open(EXPORT_RECORDS_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def add_export_record(project_name, files, sub_folder):
    """添加一条导出记录
    Args:
        project_name: 项目名称
        files: 导出的文件列表
        sub_folder: 子文件夹名称（硬件素材 或 演示）
    """
    now = datetime.now()
    export_time = now.strftime('%Y-%m-%d %H:%M:%S')
    record_id = now.strftime('%Y%m%d_%H%M%S')

    record = {
        "id": record_id,
        "project_name": project_name,
        "export_time": export_time,
        "timestamp": int(now.timestamp()),
        "files": files,
        "file_count": len(files),
        "sub_folder": sub_folder
    }

    records_data = load_export_records()
    # 在列表开头插入最新的记录（最新的在前面）
    records_data["records"].insert(0, record)
    save_export_records(records_data)


def read_xls():
    """读取总表.xls（支持 .xls 和 .xlsx 格式），返回项目列表

    Excel 列结构（当前）:
    A列: 项目名称（系统名）
    B列: 元器件清单
    C列: 功能清单
    D列: 硬件成本
    E列: 状态（可选）
    F列: 备注（可选）
    G列: 项目路径（可选）
    """
    if not os.path.exists(XLS_PATH):
        return []

    projects = []

    # 根据文件格式选择合适的库
    try:
        if XLS_PATH.endswith('.xlsx'):
            # 使用 openpyxl 读取 .xlsx
            wb = load_workbook(XLS_PATH, data_only=True)
            ws = wb.active

            for i in range(2, ws.max_row + 1):  # 从第2行开始（第1行是表头）
                name = ws.cell(i, 1).value
                if not name:
                    continue
                name = str(name).strip()

                # 读取硬件成本（D列，索引4）
                hw_cost = None
                cost_val = ws.cell(i, 4).value
                if cost_val:
                    try:
                        hw_cost = float(cost_val)
                    except (ValueError, TypeError):
                        pass

                components_raw = ws.cell(i, 2).value or ''
                functions_raw = ws.cell(i, 3).value or ''

                components_raw = str(components_raw).strip() if components_raw else ''
                functions_raw = str(functions_raw).strip() if functions_raw else ''

                projects.append({
                    'id': None,
                    'name': name,
                    'components': components_raw,
                    'functions': functions_raw,
                    'hw_cost': hw_cost,
                    'xls_row': i,
                })
        else:
            # 使用 xlrd 读取 .xls
            wb = xlrd.open_workbook(XLS_PATH, formatting_info=False)
            ws = wb.sheets()[0]

            for i in range(1, ws.nrows):
                row = ws.row_values(i)
                name = str(row[0]).strip() if row[0] else ''
                if not name:
                    continue

                # 读取硬件成本（D列，索引3）
                hw_cost = None
                if len(row) > 3 and row[3]:
                    try:
                        hw_cost = float(row[3])
                    except (ValueError, TypeError):
                        pass

                components_raw = str(row[1]).strip() if len(row) > 1 and row[1] else ''
                functions_raw = str(row[2]).strip() if len(row) > 2 and row[2] else ''

                projects.append({
                    'id': None,
                    'name': name,
                    'components': components_raw,
                    'functions': functions_raw,
                    'hw_cost': hw_cost,
                    'xls_row': i,
                })

        return projects
    except Exception as e:
        print(f'读取 Excel 失败: {e}', file=sys.stderr)
        import traceback
        traceback.print_exc()
        return []


def parse_component_name(raw: str) -> str:
    """从 'DHT11(PA4)' 或 'DHT11(1)' 提取器件名 'DHT11'"""
    return re.sub(r'\(.*?\)', '', raw).strip()


# ── 器件名归一化规则 ──────────────────────────────────────────
# 格式: (关键词列表, 规范名)
# 匹配规则: 器件名（大小写不敏感）包含任意关键词则归一化为规范名
_PART_NORMALIZE_RULES = [
    (['stm32f103c8t6'], 'STM32F103C8T6'),
    (['stm32f103'], 'STM32F103C8T6'),
    (['dht11'], 'DHT11'),
    (['dht22', 'am2302'], 'DHT22'),
    (['ds18b20'], 'DS18B20'),
    (['ds1302'], 'DS1302'),
    (['ds1307'], 'DS1307'),
    (['bmp280'], 'BMP280'),
    (['bme280'], 'BME280'),
    (['mpu6050'], 'MPU6050'),
    (['hx711'], 'HX711'),
    (['hc-sr04', 'hcsr04'], 'HC-SR04'),
    (['hc-sr501', 'hcsr501'], 'HC-SR501'),
    (['max30102'], 'MAX30102'),
    (['max6675'], 'MAX6675'),
    (['mq-2', 'mq2'], 'MQ-2'),
    (['mq-5', 'mq5'], 'MQ-5'),
    (['mq-135', 'mq135'], 'MQ-135'),
    (['neo-6m', 'neo6m', 'gps'], 'NEO-6M GPS'),
    (['bc26', 'nb-iot', 'nbiot'], 'BC26 NB-IoT'),
    (['sim800', 'gsm', 'gprs'], 'SIM800C'),
    (['air724', 'ec800', '4g模块'], '4G模块'),
    (['esp8266'], 'ESP8266'),
    (['esp32'], 'ESP32'),
    (['nrf24'], 'NRF24L01'),
    (['rc522'], 'RC522'),
    (['ssd1306', 'oled'], 'OLED SSD1306'),
    (['lcd1602', 'lcd 1602'], 'LCD1602'),
    (['st7735'], 'ST7735'),
    (['ili9341'], 'ILI9341'),
    (['pcf8591'], 'PCF8591'),
    (['tm1637'], 'TM1637'),
    (['dfplayer', 'mp3模块', 'mp3ģ'], 'DFPlayer Mini'),
    (['ld3320', '语音识别'], 'LD3320'),
    (['28byj-48', '步进电机'], '28BYJ-48步进电机'),
    (['uln2003'], 'ULN2003'),
    (['l298'], 'L298N'),
    (['继电器'], '继电器'),
    (['蜂鸣器'], '有源蜂鸣器'),
    (['舵机', 'servo'], '舵机'),
    (['按键'], '按键x4'),
    (['风扇', 'fan'], '直流风扇'),
    (['led'], 'LED'),
    (['microsd', 'sd卡', 'sd模块'], 'MicroSD卡'),
    (['光敏'], '光敏电阻'),
    (['声音传感'], '声音传感器'),
    (['火焰传感'], '火焰传感器'),
    (['土壤湿度', '电容式土壤'], '土壤湿度传感器'),
    (['扬声器'], '扬声器'),
]

def normalize_part_name(name: str) -> str:
    """将各种写法的器件名归一化为规范名；匹配不到则返回原名"""
    low = name.lower()
    for keywords, canonical in _PART_NORMALIZE_RULES:
        for kw in keywords:
            if kw.lower() in low:
                return canonical
    return name


def parse_qty_from_name(raw: str) -> int:
    """从括号里提取数量，只有纯数字才算数量，引脚不算"""
    m = re.search(r'\((\d+)\)', raw)
    if m:
        return int(m.group(1))
    return 1


def get_active_statuses():
    """未完成状态"""
    return {'待绘制', '待打板', '等板', '待交付'}


def compute_demand(xls_projects, biz_data):
    """
    统计未完成项目的器件需求
    返回: { "DHT11": { "total": 3, "projects": ["项目A(1)", "项目B(2)"] } }
    """
    active = get_active_statuses()
    demand = {}
    for p in xls_projects:
        name = p['name']
        status = biz_data.get(name, {}).get('status', '待打板')
        if status not in active:
            continue
        for line in p['components'].split('\n'):
            line = line.strip()
            if not line:
                continue
            part = normalize_part_name(parse_component_name(line))
            qty = parse_qty_from_name(line)
            if not part:
                continue
            if part not in demand:
                demand[part] = {'total': 0, 'projects': []}
            demand[part]['total'] += qty
            demand[part]['projects'].append(f"{name}({qty})")
    return demand


def get_project_folder(name: str):
    folder = os.path.join(PROJECTS_DIR, name)
    return folder if os.path.isdir(folder) else None


def get_project_files(name: str):
    folder = get_project_folder(name)
    if not folder:
        return []
    result = []
    # 根目录文件
    for f in os.listdir(folder):
        if os.path.isfile(os.path.join(folder, f)):
            result.append(f)
    # 扫描两个固定子文件夹
    for sub in ['硬件素材', '演示']:
        sub_path = os.path.join(folder, sub)
        if os.path.isdir(sub_path):
            for f in os.listdir(sub_path):
                if os.path.isfile(os.path.join(sub_path, f)):
                    result.append(f'{sub}/{f}')
    return result


# ─── API ────────────────────────────────────────────────────

@app.route('/')
def index():
    resp = make_response(render_template('index.html'))
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp


@app.route('/api/projects', methods=['GET'])
def api_get_projects():
    try:
        xls_projects = read_xls()
        biz_data = load_json()
        result = []

        # 建立 Excel 项目映射
        xls_map = {p['name']: p for p in xls_projects}

        # 遍历 JSON 中的所有项目（主数据源）
        for name, biz in biz_data.items():
            p = xls_map.get(name, {})
            folder = get_project_folder(name)
            files = get_project_files(name) if folder else []
            result.append({
                'id': p.get('id'),
                'name': name,
                'components': p.get('components', ''),
                'functions': p.get('functions', ''),
                'hw_cost': p.get('hw_cost'),
                'folder': folder,
                'files': files,
                'deliverable': biz.get('deliverable', ''),
                'price': biz.get('price', ''),
                'prepay': biz.get('prepay', ''),
                'date': biz.get('date', ''),
                'status': biz.get('status', '待绘制'),
                'note': biz.get('note', ''),
            })

        # 添加 Excel 中有但 JSON 中没有的项目
        for p in xls_projects:
            if p['name'] not in biz_data:
                folder = get_project_folder(p['name'])
                files = get_project_files(p['name']) if folder else []
                result.append({
                    'id': p.get('id'),
                    'name': p['name'],
                    'components': p.get('components', ''),
                    'functions': p.get('functions', ''),
                    'hw_cost': p.get('hw_cost'),
                    'folder': folder,
                    'files': files,
                    'deliverable': '',
                    'price': '',
                    'prepay': '',
                    'date': '',
                    'status': '待绘制',
                    'note': '',
                })

        return jsonify(result)
    except Exception as e:
        print(f"[ERROR] api_get_projects: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/projects/<path:name>', methods=['PUT'])
def api_update_project(name):
    biz_data = load_json()
    body = request.get_json(force=True)
    allowed = {'deliverable', 'price', 'prepay', 'date', 'status', 'note'}
    entry = biz_data.get(name, {})
    for k, v in body.items():
        if k in allowed:
            entry[k] = v
    biz_data[name] = entry
    save_json(biz_data)
    return jsonify({'ok': True})


@app.route('/api/stats', methods=['GET'])
def api_stats():
    biz_data = load_json()
    total_price = 0
    total_prepay = 0
    status_count = {}
    for v in biz_data.values():
        try:
            total_price += float(str(v.get('price', 0)).replace(',', '') or 0)
        except Exception:
            pass
        try:
            total_prepay += float(str(v.get('prepay', 0)).replace(',', '') or 0)
        except Exception:
            pass
        s = v.get('status', '待打板')
        status_count[s] = status_count.get(s, 0) + 1
    xls_count = len(read_xls())
    return jsonify({
        'total': xls_count,
        'total_price': total_price,
        'total_prepay': total_prepay,
        'status_count': status_count,
    })


@app.route('/api/stats/daily', methods=['GET'])
def api_stats_daily():
    """
    按日期聚合统计：每天的单子数量、总价、预付款
    返回按日期排序的列表，方便前端展示
    """
    biz_data = load_json()
    daily = {}  # { "2024-01-01": { count, price, prepay } }
    for name, v in biz_data.items():
        date = v.get('date', '')
        if not date:
            date = '未填日期'
        if date not in daily:
            daily[date] = {'date': date, 'count': 0, 'price': 0.0, 'prepay': 0.0, 'projects': []}
        daily[date]['count'] += 1
        try:
            daily[date]['price'] += float(str(v.get('price', 0)).replace(',', '') or 0)
        except Exception:
            pass
        try:
            daily[date]['prepay'] += float(str(v.get('prepay', 0)).replace(',', '') or 0)
        except Exception:
            pass
        daily[date]['projects'].append({
            'name': name,
            'price': v.get('price', ''),
            'prepay': v.get('prepay', ''),
            'status': v.get('status', '待绘制'),
        })
    # 按日期倒序排列，未填日期放最后
    result = sorted(daily.values(), key=lambda x: (x['date'] == '未填日期', x['date']), reverse=False)
    result = sorted(result, key=lambda x: (x['date'] == '未填日期', x['date']), reverse=True)
    return jsonify(result)


def _calc_estimate(project, parts_db):
    """计算单个项目的硬件成本，复用于单项和批量接口"""
    items = []
    missing = []
    total_cost = 0.0
    for line in project['components'].split('\n'):
        line = line.strip()
        if not line:
            continue
        raw_name = parse_component_name(line)
        part = normalize_part_name(raw_name)
        qty = parse_qty_from_name(line)
        if not part:
            continue
        entry = parts_db.get(part, {})
        unit_price = None
        for record in reversed(entry.get('purchases', [])):
            if record.get('unit_price') is not None:
                unit_price = float(record['unit_price'])
                break
        subtotal = round(unit_price * qty, 2) if unit_price is not None else None
        if subtotal is not None:
            total_cost += subtotal
        items.append({'part': part, 'qty': qty, 'unit_price': unit_price, 'subtotal': subtotal})
        if unit_price is None:
            missing.append(part)
    return {'ok': True, 'items': items, 'total_cost': round(total_cost, 2), 'missing': missing}


@app.route('/api/estimate/all', methods=['GET'])
def api_estimate_all():
    """批量返回所有项目的硬件成本估算，一次请求解决"""
    xls_projects = read_xls()
    parts_db = load_parts()
    result = {}
    for p in xls_projects:
        result[p['name']] = _calc_estimate(p, parts_db)
    return jsonify(result)


@app.route('/api/estimate/<path:name>', methods=['GET'])
def api_estimate(name):
    """
    预报价：根据项目元器件清单 + 器件库单价，估算硬件成本
    返回: { items: [{part, qty, unit_price, subtotal}], total_cost, missing: [无价格的器件] }
    """
    xls_projects = read_xls()
    parts_db = load_parts()
    project = next((p for p in xls_projects if p['name'] == name), None)
    if not project:
        return jsonify({'ok': False, 'error': '项目不存在'}), 404
    return jsonify(_calc_estimate(project, parts_db))


def write_hw_costs_to_xls(costs: dict):
    """
    将硬件成本写入总表 D 列（index 3）。
    costs: { "项目名": float 或 None }
    """
    lock_path = XLS_PATH + '.lock'
    while True:
        try:
            lf = open(lock_path, 'x')
            break
        except FileExistsError:
            time.sleep(0.3)

    updated = 0
    try:
        rb = xlrd.open_workbook(XLS_PATH, formatting_info=True)
        wb = xl_copy(rb)
        ws = wb.get_sheet(0)
        sheet = rb.sheets()[0]
        for i in range(1, sheet.nrows):
            row = sheet.row_values(i)
            name = str(row[0]).strip() if row[0] else ''
            if name in costs and costs[name] is not None:
                ws.write(i, 3, round(float(costs[name]), 2))
                updated += 1
        wb.save(XLS_PATH)
    finally:
        lf.close()
        os.remove(lock_path)
    return updated


@app.route('/api/costs/update', methods=['POST'])
def api_costs_update():
    """
    一键更新所有项目的硬件成本：计算后写入总表 D 列
    """
    xls_projects = read_xls()
    parts_db = load_parts()

    costs = {}
    for p in xls_projects:
        result = _calc_estimate(p, parts_db)
        costs[p['name']] = result['total_cost'] if result['total_cost'] > 0 else None

    try:
        updated = write_hw_costs_to_xls(costs)
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

    return jsonify({'ok': True, 'updated': updated, 'total': len(xls_projects)})


@app.route('/api/open_folder/<path:name>', methods=['POST'])
def api_open_folder(name):
    folder = get_project_folder(name)
    if folder and os.path.isdir(folder):
        os.startfile(folder)
        return jsonify({'ok': True})
    return jsonify({'ok': False, 'error': '文件夹不存在'}), 404


# ─── 文件上传配置 ──────────────────────────────────────────────
_UPLOAD_FOLDER_MAP = {
    0: '硬件素材',   # 待绘制
    1: '硬件素材',   # 待打板
    2: '演示',       # 等板
    3: '演示',       # 待交付
    4: '演示',       # 已完成
}
_STATUS_WEIGHT = {
    '待绘制': 0,
    '待打板': 1,
    '等板': 2,
    '待交付': 3,
    '已完成': 4,
}
_ALLOWED_EXTENSIONS = {
    '.zip', '.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp',
    '.mp4', '.mov', '.avi', '.mkv',
    '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.md', '.txt', '.csv',
    '.sch', '.pcb', '.dwg', '.dxf', '.step', '.stl', '.epro2',
}


# ─── 文件自动重命名规则 ───────────────────────────────────────────
def get_auto_renamed_file(filename, project_name, sub_folder_name):
    """
    根据项目阶段和文件类型自动重命名文件

    硬件素材阶段（待绘制、待打板）:
      - .zip → 项目名+PCB制板文件
      - .epro2 → 项目名+eda工程
      - .jpg/.jpeg/.png/.gif/.bmp → 项目名+图片展示
      - .pdf → 项目名+原理图
      - 其他 → 保持原名

    演示阶段（等板、待交付、已完成）:
      - .jpg/.jpeg/.png/.gif/.bmp → 项目名_展示_01/02/03... (需要检查重复)
      - .mp4/.mov/.avi/.mkv → 项目名+功能演示
      - 其他 → 保持原名
    """
    base, ext = os.path.splitext(filename)
    ext_lower = ext.lower()

    if sub_folder_name == '硬件素材':
        # 硬件素材阶段的重命名规则
        if ext_lower == '.zip':
            return f'{project_name}+PCB制板文件{ext}'
        elif ext_lower == '.epro2':
            return f'{project_name}+eda工程{ext}'
        elif ext_lower in {'.jpg', '.jpeg', '.png', '.gif', '.bmp'}:
            return f'{project_name}+图片展示{ext}'
        elif ext_lower == '.pdf':
            return f'{project_name}+原理图{ext}'
        else:
            # 其他格式保持原名
            return filename

    elif sub_folder_name == '演示':
        # 演示阶段的重命名规则
        if ext_lower in {'.jpg', '.jpeg', '.png', '.gif', '.bmp'}:
            # 图片需要编号，返回特殊格式，后续处理
            return ('__IMAGE__', project_name, ext)
        elif ext_lower in {'.mp4', '.mov', '.avi', '.mkv'}:
            return f'{project_name}+功能演示{ext}'
        else:
            # 其他格式保持原名
            return filename

    return filename


def get_numbered_image_name(project_name, target_dir, ext):
    """
    为演示阶段的图片生成带序号的名称 (项目名_展示_01/02/03...)
    通过扫描目录中已有的同类图片来确定下一个序号
    注：所有图片（不管扩展名）共享同一个编号序列
    """
    pattern = f"{project_name}_展示_"
    counter = 1

    # 扫描目录中所有匹配的图片名称，找到最大的序号
    if os.path.isdir(target_dir):
        for filename in os.listdir(target_dir):
            # 检查是否匹配 "项目名_展示_XX" 的格式
            if filename.startswith(pattern):
                # 尝试提取序号
                basename = filename[len(pattern):]  # 去掉前缀
                # 提取数字部分（可能是 "01.jpg", "01_backup.jpg" 等）
                match = re.match(r'^(\d+)', basename)
                if match:
                    try:
                        num = int(match.group(1))
                        counter = max(counter, num + 1)
                    except ValueError:
                        pass

    # 找到下一个未使用的序号
    while True:
        numbered_name = f'{pattern}{counter:02d}{ext}'
        full_path = os.path.join(target_dir, numbered_name)
        if not os.path.exists(full_path):
            return numbered_name
        counter += 1


@app.route('/api/projects/<path:name>/upload', methods=['POST'])
def api_upload_file(name):
    """
    文件上传接口
    - 自动根据项目当前阶段决定存入 硬件素材/ 或 演示/ 子文件夹
    - 根据文件类型自动重命名（硬件素材和演示阶段有不同规则）
    - 支持 multipart/form-data 多文件上传
    - 返回: { ok, saved: ["硬件素材/xxx.pdf", ...], errors: [...] }
    """
    folder = get_project_folder(name)
    if not folder:
        return jsonify({'ok': False, 'error': '项目文件夹不存在'}), 404

    # 首先检查前端是否指定了上传目录
    upload_dir = request.form.get('upload_dir', '').strip()

    # 如果前端指定了目录，使用前端的选择；否则根据项目状态自动选择
    if upload_dir in ['硬件', '演示']:
        sub_dir_name = '硬件素材' if upload_dir == '硬件' else '演示'
    else:
        # 自动选择：根据项目当前状态推断目标子文件夹
        biz_data = load_json()
        status = biz_data.get(name, {}).get('status', '待绘制')
        weight = _STATUS_WEIGHT.get(status, 0)
        sub_dir_name = _UPLOAD_FOLDER_MAP.get(weight, '硬件素材')

    target_dir = os.path.join(folder, sub_dir_name)

    # 子文件夹不存在则自动创建
    os.makedirs(target_dir, exist_ok=True)

    files = request.files.getlist('files')
    if not files:
        return jsonify({'ok': False, 'error': '未收到文件'}), 400

    saved = []
    errors = []
    for f in files:
        original_name = f.filename
        if not original_name:
            continue
        # 安全检查：只允许白名单扩展名
        ext = os.path.splitext(original_name)[1].lower()
        if ext not in _ALLOWED_EXTENSIONS:
            errors.append(f'{original_name}（不支持的文件类型）')
            continue

        # 文件名安全处理：去除路径分隔符，防止路径穿越
        safe_name = os.path.basename(original_name)

        # 自动重命名：根据阶段和文件类型应用重命名规则
        renamed = get_auto_renamed_file(safe_name, name, sub_dir_name)

        # 处理演示阶段图片的编号情况
        if isinstance(renamed, tuple) and renamed[0] == '__IMAGE__':
            # 这是演示阶段的图片，需要生成编号
            _, proj_name, ext_part = renamed
            safe_name = get_numbered_image_name(proj_name, target_dir, ext_part)
        else:
            safe_name = renamed

        dest_path = os.path.join(target_dir, safe_name)

        # 若同名文件已存在，自动重命名（加数字后缀）
        if os.path.exists(dest_path):
            base, ext_part = os.path.splitext(safe_name)
            i = 1
            while os.path.exists(dest_path):
                safe_name = f'{base}_{i}{ext_part}'
                dest_path = os.path.join(target_dir, safe_name)
                i += 1

        f.save(dest_path)
        saved.append(f'{sub_dir_name}/{safe_name}')

    return jsonify({
        'ok': True,
        'sub_folder': sub_dir_name,
        'saved': saved,
        'errors': errors,
    })


# ─── 器件库 API ──────────────────────────────────────────────

@app.route('/api/parts', methods=['GET'])
def api_get_parts():
    """
    返回器件库全量数据，附带需求计算
    每条: { name, stock, demand, diff, projects, purchases }
    """
    xls_projects = read_xls()
    biz_data = load_json()
    parts_db = load_parts()
    demand = compute_demand(xls_projects, biz_data)

    # 合并：器件库已有 + 项目里出现过的都纳入
    all_parts = set(parts_db.keys()) | set(demand.keys())

    result = []
    for part in sorted(all_parts):
        entry = parts_db.get(part, {})
        stock = entry.get('stock', 0)
        d = demand.get(part, {}).get('total', 0)
        projects = demand.get(part, {}).get('projects', [])
        result.append({
            'name': part,
            'stock': stock,
            'demand': d,
            'diff': stock - d,          # 正=富余，负=缺货
            'projects': projects,        # 哪些项目需要
            'purchases': entry.get('purchases', []),
        })

    # 缺货的排前面
    result.sort(key=lambda x: (x['diff'] >= 0, x['name']))
    return jsonify(result)


@app.route('/api/parts/sync', methods=['POST'])
def api_parts_sync():
    """从项目列表同步器件名到器件库（新器件默认库存0）"""
    xls_projects = read_xls()
    biz_data = load_json()
    parts_db = load_parts()
    demand = compute_demand(xls_projects, biz_data)
    added = 0
    for part in demand:
        if part not in parts_db:
            parts_db[part] = {'stock': 0, 'purchases': []}
            added += 1
    save_parts(parts_db)
    return jsonify({'ok': True, 'added': added, 'total': len(parts_db)})


@app.route('/api/parts/<path:name>/stock', methods=['PUT'])
def api_set_stock(name):
    """直接设置某器件的库存数量"""
    parts_db = load_parts()
    body = request.get_json(force=True)
    entry = parts_db.get(name, {'stock': 0, 'purchases': []})
    entry['stock'] = int(body.get('stock', 0))
    parts_db[name] = entry
    save_parts(parts_db)
    return jsonify({'ok': True, 'stock': entry['stock']})


@app.route('/api/parts/<path:name>/purchase', methods=['POST'])
def api_purchase(name):
    """采购入库：增加库存 + 记录采购记录"""
    parts_db = load_parts()
    body = request.get_json(force=True)
    qty = int(body.get('qty', 0))
    note = str(body.get('note', '')).strip()
    unit_price = body.get('unit_price')  # 可为 null
    if unit_price is not None:
        try:
            unit_price = float(unit_price)
        except (ValueError, TypeError):
            unit_price = None
    if qty <= 0:
        return jsonify({'ok': False, 'error': '数量必须大于0'}), 400

    entry = parts_db.get(name, {'stock': 0, 'purchases': []})
    entry['stock'] = entry.get('stock', 0) + qty
    if 'purchases' not in entry:
        entry['purchases'] = []
    record = {
        'qty': qty,
        'date': datetime.now().strftime('%Y-%m-%d'),
        'note': note,
    }
    if unit_price is not None:
        record['unit_price'] = unit_price
    entry['purchases'].append(record)
    parts_db[name] = entry
    save_parts(parts_db)
    return jsonify({'ok': True, 'stock': entry['stock']})


# ─── 导出相关 API ────────────────────────────────────────────────

@app.route('/api/projects/<path:project_name>/export-files', methods=['POST'])
def api_export_files(project_name):
    """
    导出文件为 ZIP 压缩包
    Request JSON: { "files": ["文件1.pdf", "文件2.zip", ...] }
    Response: 返回 ZIP 文件下载 + 记录到导出历史
    """
    folder = get_project_folder(project_name)
    if not folder:
        return jsonify({'ok': False, 'error': '项目文件夹不存在'}), 404

    data = request.get_json() or {}
    files_to_export = data.get('files', [])

    if not files_to_export:
        return jsonify({'ok': False, 'error': '请至少选择一个文件'}), 400

    # 读取项目当前阶段来确定文件来源子文件夹
    biz_data = load_json()
    status = biz_data.get(project_name, {}).get('status', '待绘制')
    weight = _STATUS_WEIGHT.get(status, 0)
    sub_dir_name = _UPLOAD_FOLDER_MAP.get(weight, '硬件素材')
    source_dir = os.path.join(folder, sub_dir_name)

    # 创建 ZIP 文件
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for filename in files_to_export:
            # 安全检查：防止路径穿越
            safe_filename = os.path.basename(filename)
            file_path = os.path.join(source_dir, safe_filename)

            # 验证文件存在且在允许的目录内
            if not os.path.exists(file_path) or not os.path.abspath(file_path).startswith(os.path.abspath(source_dir)):
                continue  # 跳过不存在或非法路径的文件

            # 添加文件到 ZIP
            arcname = safe_filename  # ZIP 内的文件名
            zf.write(file_path, arcname=arcname)

    # 生成压缩包名称（包含时间戳）
    now = datetime.now()
    timestamp = now.strftime('%Y%m%d_%H%M%S')
    zip_filename = f'{project_name}_{sub_dir_name}_{timestamp}.zip'

    # 记录导出操作
    add_export_record(project_name, files_to_export, sub_dir_name)

    # 返回 ZIP 文件
    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=zip_filename
    )


@app.route('/api/export-history', methods=['GET'])
def api_get_export_history():
    """
    获取所有导出记录
    Response: { "ok": true, "records": [...] }
    """
    records_data = load_export_records()
    return jsonify({
        'ok': True,
        'records': records_data.get('records', [])
    })


@app.route('/api/export-history/<record_id>', methods=['DELETE'])
def api_delete_export_record(record_id):
    """
    删除一条导出记录
    Response: { "ok": true, "message": "记录已删除" }
    """
    records_data = load_export_records()
    records = records_data.get('records', [])

    # 找到并删除匹配的记录
    for i, record in enumerate(records):
        if record.get('id') == record_id:
            records.pop(i)
            save_export_records(records_data)
            return jsonify({'ok': True, 'message': '记录已删除'})

    return jsonify({'ok': False, 'error': '记录不存在'}), 404


# ─── 通用导出系统（Universal Export System） ──────────────

class BaseExporter:
    """导出器基类"""

    def __init__(self, projects_data, all_projects):
        """
        初始化导出器
        Args:
            projects_data: 商务信息字典 (来自 load_json)
            all_projects: 项目列表 (来自 read_xls)
        """
        self.projects_data = projects_data
        self.all_projects = all_projects

    def export(self, project_names):
        """
        导出文件
        Args:
            project_names: 要导出的项目名称列表

        Returns:
            (file_buffer, filename) 元组
        """
        raise NotImplementedError("子类必须实现 export() 方法")

    def get_filename(self, project_names):
        """生成文件名"""
        raise NotImplementedError("子类必须实现 get_filename() 方法")

    def _get_project_by_name(self, name):
        """根据项目名获取项目对象"""
        for p in self.all_projects:
            if p.get('name') == name:
                return p
        return None


class PCBExporter(BaseExporter):
    """PCB制板文件导出器"""

    def export(self, project_names):
        """
        导出PCB文件为ZIP
        """
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for project_name in project_names:
                project = self._get_project_by_name(project_name)
                if not project:
                    continue

                # 获取项目文件夹
                folder = get_project_folder(project_name)
                if not folder:
                    continue

                # 获取PCB文件（直接扫描文件夹，不依赖 read_xls 的 files 字段）
                for file_path in get_project_files(project_name):
                    if '+PCB制板文件.zip' in file_path:
                        # file_path 可能是相对路径，如 "硬件素材/LED控制器+PCB制板文件.zip"
                        full_path = os.path.join(folder, file_path)
                        if os.path.exists(full_path):
                            zf.write(full_path, arcname=os.path.basename(file_path))

        zip_buffer.seek(0)
        return (zip_buffer, self.get_filename(project_names))

    def get_filename(self, project_names):
        """生成PCB导出文件名"""
        now = datetime.now()
        timestamp = now.strftime('%Y%m%d_%H%M%S')
        if len(project_names) == 1:
            return f'{project_names[0]}_PCB文件_{timestamp}.zip'
        else:
            return f'PCB文件_{timestamp}.zip'


class DocxExporter(BaseExporter):
    """DOCX标签导出器"""

    def export(self, project_names):
        """
        导出项目标签为DOCX
        """
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # 添加标题
        title = doc.add_heading('项目分类标签', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加生成时间
        timestamp_para = doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        timestamp_para.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').font.size = Pt(9)

        doc.add_paragraph()  # 添加空行

        # 遍历项目并添加到文档
        for idx, project_name in enumerate(project_names, 1):
            project = self._get_project_by_name(project_name)
            if not project:
                continue

            # 获取项目编号（从 XLS 读取的编号）
            project_id = project.get('id') or idx  # id 为 None 时使用序号

            # 创建项目标签表格（1行3列：编号、项目信息、功能）
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            # 设置列宽
            table.columns[0].width = Inches(0.8)
            table.columns[1].width = Inches(2.5)
            table.columns[2].width = Inches(2.5)

            # 编号单元格 - 使用从 Excel 读取的编号
            cell_num = table.rows[0].cells[0]
            para_num = cell_num.paragraphs[0]
            para_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_num = para_num.add_run(f'[{project_id:03d}]')
            run_num.font.size = Pt(14)
            run_num.font.bold = True

            # 项目信息单元格（左侧）
            cell_info = table.rows[0].cells[1]
            cell_info.vertical_alignment = 1  # 顶部对齐
            # 清除默认段落
            cell_info.paragraphs[0].text = ''

            # 项目名称
            para_name = cell_info.paragraphs[0]
            run_name = para_name.add_run(project_name)
            run_name.font.size = Pt(12)
            run_name.font.bold = True

            # 元器件列表
            components_text = project.get('components', '').strip()
            if components_text:
                para_comp = cell_info.add_paragraph()
                run_comp_label = para_comp.add_run('元器件：')
                run_comp_label.font.size = Pt(9)
                run_comp_label.font.bold = True
                run_comp_text = para_comp.add_run(components_text)
                run_comp_text.font.size = Pt(9)

            # 功能单元格（右侧）
            cell_func = table.rows[0].cells[2]
            cell_func.vertical_alignment = 1  # 顶部对齐
            cell_func.paragraphs[0].text = ''

            para_func = cell_func.paragraphs[0]
            run_func_label = para_func.add_run('功能：')
            run_func_label.font.size = Pt(9)
            run_func_label.font.bold = True
            run_func_text = para_func.add_run(project.get('functions', '').strip())
            run_func_text.font.size = Pt(9)

            # 添加分隔线
            if idx < len(project_names):
                doc.add_paragraph()

        # 保存到字节缓冲区
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return (docx_buffer, self.get_filename(project_names))

    def get_filename(self, project_names):
        """生成DOCX导出文件名"""
        now = datetime.now()
        timestamp = now.strftime('%Y%m%d_%H%M%S')
        return f'项目标签_{timestamp}.docx'


class ExportFactory:
    """导出工厂类"""

    @staticmethod
    def create_exporter(export_type, projects_data, all_projects):
        """
        创建导出器实例

        Args:
            export_type: 导出类型 ('pcb' 或 'docx')
            projects_data: 商务信息数据
            all_projects: 项目列表

        Returns:
            导出器实例
        """
        if export_type == 'pcb':
            return PCBExporter(projects_data, all_projects)
        elif export_type == 'docx':
            return DocxExporter(projects_data, all_projects)
        else:
            raise ValueError(f'不支持的导出类型: {export_type}')


@app.route('/api/export', methods=['POST'])
def api_universal_export():
    """
    通用导出API

    Request JSON:
    {
        "projects": ["项目1", "项目2"],
        "type": "pcb" | "docx",
        "stages": ["等板", "待交付"]
    }

    Response: 返回导出文件或错误信息
    """
    data = request.get_json() or {}

    # 验证请求数据
    project_names = data.get('projects', [])
    export_type = data.get('type', 'pcb')
    stages = data.get('stages', [])

    if not isinstance(project_names, list) or len(project_names) == 0:
        return jsonify({'ok': False, 'error': '请提供有效的项目列表'}), 400

    if export_type not in ['pcb', 'docx']:
        return jsonify({'ok': False, 'error': '不支持的导出类型'}), 400

    try:
        # 加载所有数据
        projects_data = load_json()
        all_projects = read_xls()

        # 创建导出器
        exporter = ExportFactory.create_exporter(export_type, projects_data, all_projects)

        # 执行导出
        file_buffer, filename = exporter.export(project_names)

        # 记录导出操作
        record = {
            "id": datetime.now().strftime('%Y%m%d_%H%M%S'),
            "export_type": export_type,
            "projects": project_names,
            "stages": stages,
            "export_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "timestamp": int(datetime.now().timestamp()),
            "file_count": len(project_names)
        }
        records_data = load_export_records()
        records_data["records"].insert(0, record)
        save_export_records(records_data)

        # 返回文件
        return send_file(
            file_buffer,
            mimetype='application/zip' if export_type == 'pcb' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# 不在此处启动 Flask app，由 app_main.py 统一启动
